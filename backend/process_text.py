import os
import json
import PyPDF2
import docx
from fpdf import FPDF
import google.generativeai as genai
import os
from dotenv import load_dotenv  # Import the load_dotenv function

# Load environment variables from .env file
load_dotenv()

# Load Gemini API Key from Environment Variable (More Secure)
API_KEY = os.getenv("GEMINI_API_KEY")  # No default value for security
if API_KEY is None:
    raise ValueError("API key not found. Please set the GEMINI_API_KEY environment variable.")

genai.configure(api_key=API_KEY)

# # Load Gemini API Key from Environment Variable (More Secure)
# API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyA46xUiYFotA4XeUlJ5LEtPte4EDKkL-LQ")  
# genai.configure(api_key=API_KEY)

# Directory for storing output files
OUTPUT_FOLDER = "generated_files"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def extract_text(file_path):
    """Extract text from PDF, DOCX, or TXT files."""
    if file_path.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith(".docx"):
        return extract_text_from_docx(file_path)
    elif file_path.endswith(".txt"):
        return extract_text_from_txt(file_path)
    else:
        raise ValueError("Unsupported file format")

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    try:
        text = ""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs]).strip()
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

def extract_text_from_txt(file_path):
    """Extract text from a TXT file."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except Exception as e:
        print(f"Error extracting text from TXT: {e}")
        return ""

def read_user_input(config_file):
    """Read the number of questions and difficulty level from JSON config file."""
    try:
        with open(config_file, "r", encoding="utf-8") as f:
            data = json.load(f)

        num_questions = data.get("num_questions")

        try:
            num_questions = int(num_questions) if num_questions else 5  # Ensure integer conversion
        except (ValueError, TypeError):
            num_questions = 5  # Default fallback

        difficulty = data.get("difficulty", "medium").lower()  # Default to 'medium'

        print(f"Read from config: num_questions={num_questions}, difficulty={difficulty}")  # Debugging
        return num_questions, difficulty

    except Exception as e:
        print(f"Error reading config file: {e}")
        return 5, "medium"  # Defaults if file read fails

def generate_questions(text, num_questions=5, difficulty="medium"):
    """Generate questions using Google Gemini API."""
    print(f"Generating {num_questions} {difficulty}-difficulty questions")  # Debugging
    
    if not text:
        return ["Error: No text found in the file."]
    
    # prompt = f"Generate exactly {num_questions} distinct questions with {difficulty} difficulty based on the following text. Ensure no extra or missing questions:\n{text[:2000]}"
    # prompt = f"Generate exactly {num_questions} distinct questions with {difficulty} difficulty based on the following text. Ensure no extra or missing questions:\n{text[:2000]}"
    prompt = f"""
    You are an expert researcher and university professor specializing in {difficulty} difficulty-level assessments. 
    Your task is to **generate exactly {num_questions} distinct and well-structured questions** based on the provided text.

    ### **STRICT INSTRUCTIONS:**
    1. Generate **precisely {num_questions}** questions—no more, no less.  
    2. Each question **must** have a corresponding answer.  
    3. Follow the exact structured format below:  

    ### **Output Format:**
    #### Question {num_questions} (Total Questions: {num_questions})
    - **Question:** [Clearly written question]
    - **Answer:** [Concise but accurate answer]
    - **Type:** [MCQ / Short Answer / Descriptive]
    - **Difficulty Level:** [Easy / Medium / Hard]
    - **Topic Category:** [Main topic related to the question]

    ### **Example Output (DO NOT CHANGE FORMAT):**
    #### Question 1 (Total Questions: 3)
    - **Question:** What is machine learning, and how does it differ from traditional programming?
    - **Answer:** Machine learning is a subset of AI that enables computers to learn from data without explicit programming. Unlike traditional programming, which follows predefined rules, ML models adapt based on patterns in the data.
    - **Type:** Short Answer
    - **Difficulty Level:** Medium
    - **Topic Category:** Machine Learning Basics

    #### Question 2 (Total Questions: 3)
    - **Question:** Which of the following algorithms is supervised learning?
    a) K-Means Clustering  
    b) Decision Trees  
    c) Apriori Algorithm  
    d) DBSCAN  
    - **Answer:** b) Decision Trees  
    - **Type:** Multiple Choice Question (MCQ)  
    - **Difficulty Level:** Easy  
    - **Topic Category:** Supervised Learning  

    #### Question 3 (Total Questions: 3)
    - **Question:** Explain the concept of overfitting in machine learning.
    - **Answer:** Overfitting occurs when a machine learning model learns patterns specific to training data but fails to generalize to new data. It results in high training accuracy but poor performance on test data.
    - **Type:** Descriptive
    - **Difficulty Level:** Hard
    - **Topic Category:** Model Optimization

    Now, using the following text, **generate exactly {num_questions} questions in this format**:

    ### **Text to Base Questions On:**  
    {text[:2000]}  
    """



    try:
        model = genai.GenerativeModel("gemini-1.5-pro-latest")
        response = model.generate_content(prompt)

        # Extract questions safely
        questions = []
        if hasattr(response, "text") and response.text:
            questions = [q.strip() for q in response.text.split("\n") if q.strip()]
        elif hasattr(response, "parts"):
            questions = [part.text.strip() for part in response.parts if part.text]

        print(f"Generated Questions: {questions}")  # Debugging
        return questions if questions else ["No questions generated."]
    
    except Exception as e:
        print(f"Error generating questions: {e}")
        return [f"Error generating questions: {str(e)}"]

def save_as_txt(questions):
    """Save questions to a TXT file."""
    file_path = os.path.join(OUTPUT_FOLDER, "generated_questions.txt")
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write("\n".join(questions))
        return file_path
    except Exception as e:
        print(f"Error saving TXT file: {e}")
        return None

def save_as_docx(questions):
    """Save questions to a DOCX file."""
    file_path = os.path.join(OUTPUT_FOLDER, "generated_questions.docx")
    try:
        doc = docx.Document()
        for q in questions:
            doc.add_paragraph(q)
        doc.save(file_path)
        return file_path
    except Exception as e:
        print(f"Error saving DOCX file: {e}")
        return None

def save_as_pdf(questions):
    """Save questions to a PDF file."""
    file_path = os.path.join(OUTPUT_FOLDER, "generated_questions.pdf")
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=10)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        for q in questions:
            pdf.multi_cell(0, 10, q)  # Use multi_cell for better text wrapping
        
        pdf.output(file_path)
        return file_path
    except Exception as e:
        print(f"Error saving PDF file: {e}")
        return None

# Example Test
if __name__ == "__main__":
    config_file = "questiongenerator.json"  # Ensure this file contains valid JSON
    num_questions, difficulty = read_user_input(config_file)

    sample_text = "Machine learning is a subset of AI that enables computers to learn from data and make decisions without explicit programming."
    # Debugging print statements
    print(f"DEBUG: num_questions={num_questions}, difficulty={difficulty}")

    
    questions = generate_questions(sample_text, num_questions=num_questions, difficulty=difficulty)
    
    print("Final Generated Questions:", questions)  # Debugging: Check output
